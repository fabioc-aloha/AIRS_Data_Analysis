"""Format tables in Word documents with professional styling.

Usage:
    python format_word_tables.py <input.docx> [output.docx]

Applies: Microsoft blue headers, alternating row shading, borders,
         anti-pagination (cantSplit + keepWithNext), consistent fonts.
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH


# --- Style constants ---
HEADER_BG = "0078D4"       # Microsoft blue
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
EVEN_ROW_BG = "F0F0F0"     # Light gray
ODD_ROW_BG = "FFFFFF"      # White
BORDER_OUTER = "666666"
BORDER_INNER = "AAAAAA"
HEADER_FONT_SIZE = Pt(10)
DATA_FONT_SIZE = Pt(9)
FONT_NAME = "Calibri"


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tc_pr.append(borders)

    for side, color in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if color:
            border_el = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" w:sz="4" '
                f'w:space="0" w:color="{color}"/>'
            )
            existing = borders.find(qn(f"w:{side}"))
            if existing is not None:
                borders.remove(existing)
            borders.append(border_el)


def set_row_cant_split(row):
    """Prevent row from splitting across pages."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(parse_xml(f'<w:cantSplit {nsdecls("w")} w:val="true"/>'))


def set_row_keep_with_next(row):
    """Keep this row on same page as the next row (used for header)."""
    for cell in row.cells:
        for para in cell.paragraphs:
            pPr = para._p.get_or_add_pPr()
            existing = pPr.find(qn("w:keepNext"))
            if existing is None:
                pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")} w:val="true"/>'))


def format_cell_text(cell, font_size, bold=False, color=None):
    """Apply font formatting to all runs in a cell."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = font_size
            run.font.name = FONT_NAME
            run.font.bold = bold
            if color:
                run.font.color.rgb = color
        # If no runs, create one from the text
        if not para.runs and para.text.strip():
            text = para.text
            para.clear()
            run = para.add_run(text)
            run.font.size = font_size
            run.font.name = FONT_NAME
            run.font.bold = bold
            if color:
                run.font.color.rgb = color


def format_table(table):
    """Apply professional formatting to a single table."""
    if not table.rows:
        return

    # --- Header row ---
    header_row = table.rows[0]
    set_row_cant_split(header_row)
    set_row_keep_with_next(header_row)

    for cell in header_row.cells:
        set_cell_shading(cell, HEADER_BG)
        format_cell_text(cell, HEADER_FONT_SIZE, bold=True, color=HEADER_FG)
        set_cell_borders(
            cell,
            top=BORDER_OUTER, bottom=BORDER_OUTER,
            left=BORDER_OUTER, right=BORDER_OUTER,
        )

    # --- Data rows ---
    for idx, row in enumerate(table.rows[1:], start=1):
        set_row_cant_split(row)

        bg = EVEN_ROW_BG if idx % 2 == 0 else ODD_ROW_BG
        for cell in row.cells:
            set_cell_shading(cell, bg)
            format_cell_text(cell, DATA_FONT_SIZE, bold=False)
            set_cell_borders(
                cell,
                top=BORDER_INNER, bottom=BORDER_INNER,
                left=BORDER_INNER, right=BORDER_INNER,
            )

    # Outer border on last row bottom
    if len(table.rows) > 1:
        for cell in table.rows[-1].cells:
            set_cell_borders(cell, bottom=BORDER_OUTER)


def format_document(input_path, output_path=None):
    """Open a .docx, format all tables, save."""
    doc = Document(input_path)
    table_count = len(doc.tables)

    for table in doc.tables:
        format_table(table)

    out = output_path or input_path
    doc.save(out)
    return table_count


def main():
    if len(sys.argv) < 2:
        print("Usage: python format_word_tables.py <input.docx> [output.docx]")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else input_path

    if not input_path.exists():
        print(f"Error: {input_path} not found")
        sys.exit(1)

    count = format_document(str(input_path), str(output_path))
    print(f"Formatted {count} table(s) in {output_path}")


if __name__ == "__main__":
    main()
