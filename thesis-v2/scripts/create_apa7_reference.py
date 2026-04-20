#!/usr/bin/env python3
"""Create an APA 7th Edition reference document for pandoc Word output.

This script creates a properly styled reference.docx that pandoc uses
to format Word documents according to APA 7th Edition guidelines.

Usage:
    python create_apa7_reference.py
    
Output:
    templates/reference.docx

Requirements:
    pip install python-docx lxml
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# APA 7 Constants
APA_FONT = 'Times New Roman'
APA_SIZE = 12
APA_FIRST_INDENT = 0.5  # inches
APA_MARGIN = 1.0  # inches


def set_document_defaults(doc):
    """Set document-wide default font (critical for pandoc inheritance)."""
    # Access document styles element
    styles_element = doc.styles.element
    
    # Create or get docDefaults
    doc_defaults = styles_element.find(qn('w:docDefaults'))
    if doc_defaults is None:
        doc_defaults = OxmlElement('w:docDefaults')
        styles_element.insert(0, doc_defaults)
    
    # Create rPrDefault (run properties default)
    rpr_default = doc_defaults.find(qn('w:rPrDefault'))
    if rpr_default is None:
        rpr_default = OxmlElement('w:rPrDefault')
        doc_defaults.append(rpr_default)
    
    # Create rPr (run properties)
    rpr = rpr_default.find(qn('w:rPr'))
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        rpr_default.append(rpr)
    
    # Set fonts
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:ascii'), APA_FONT)
    rfonts.set(qn('w:hAnsi'), APA_FONT)
    rfonts.set(qn('w:eastAsia'), APA_FONT)
    rfonts.set(qn('w:cs'), APA_FONT)
    rpr.append(rfonts)
    
    # Set size (in half-points, so 12pt = 24)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(APA_SIZE * 2))
    rpr.append(sz)
    
    sz_cs = OxmlElement('w:szCs')
    sz_cs.set(qn('w:val'), str(APA_SIZE * 2))
    rpr.append(sz_cs)
    
    # Create pPrDefault (paragraph properties default)
    ppr_default = doc_defaults.find(qn('w:pPrDefault'))
    if ppr_default is None:
        ppr_default = OxmlElement('w:pPrDefault')
        doc_defaults.append(ppr_default)
    
    # Create pPr (paragraph properties)
    ppr = ppr_default.find(qn('w:pPr'))
    if ppr is None:
        ppr = OxmlElement('w:pPr')
        ppr_default.append(ppr)
    
    # Set default line spacing to double (480 twips = 24pt for 12pt font)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '480')
    spacing.set(qn('w:lineRule'), 'auto')
    ppr.append(spacing)


def set_style_font(style, name=APA_FONT, size=APA_SIZE, bold=False, italic=False):
    """Set font properties for a style."""
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    # Set east-asia and complex script fonts too
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)
    
    # Also set size explicitly in XML
    sz = rpr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rpr.append(sz)
    sz.set(qn('w:val'), str(size * 2))
    
    sz_cs = rpr.find(qn('w:szCs'))
    if sz_cs is None:
        sz_cs = OxmlElement('w:szCs')
        rpr.append(sz_cs)
    sz_cs.set(qn('w:val'), str(size * 2))


def create_apa7_reference():
    """Create an APA 7 compliant reference document for pandoc."""
    doc = Document()
    
    # ============================================================
    # SET DOCUMENT-WIDE DEFAULTS (critical for pandoc)
    # ============================================================
    set_document_defaults(doc)
    
    # ============================================================
    # PAGE SETUP (APA 7: 1" margins all sides)
    # ============================================================
    for section in doc.sections:
        section.top_margin = Inches(APA_MARGIN)
        section.bottom_margin = Inches(APA_MARGIN)
        section.left_margin = Inches(APA_MARGIN)
        section.right_margin = Inches(APA_MARGIN)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)
    
    # ============================================================
    # NORMAL STYLE (Body text)
    # APA 7: Times New Roman 12pt, double-spaced, 0.5" first-line indent
    # ============================================================
    normal_style = doc.styles['Normal']
    set_style_font(normal_style, APA_FONT, APA_SIZE)
    
    normal_para = normal_style.paragraph_format
    normal_para.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal_para.space_before = Pt(0)
    normal_para.space_after = Pt(0)
    normal_para.first_line_indent = Inches(APA_FIRST_INDENT)
    normal_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal_para.widow_control = True
    
    # ============================================================
    # BODY TEXT STYLE (pandoc uses this for paragraphs)
    # ============================================================
    try:
        body_text = doc.styles.add_style('Body Text', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        body_text = doc.styles['Body Text']
    
    body_text.base_style = normal_style
    set_style_font(body_text, APA_FONT, APA_SIZE)
    body_text.paragraph_format.first_line_indent = Inches(APA_FIRST_INDENT)
    body_text.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    body_text.paragraph_format.space_before = Pt(0)
    body_text.paragraph_format.space_after = Pt(0)
    
    # ============================================================
    # HEADING STYLES (APA 7 has 5 levels - all same size, different formatting)
    # ============================================================
    
    # Heading 1: Centered, Bold, Title Case
    h1 = doc.styles['Heading 1']
    set_style_font(h1, APA_FONT, APA_SIZE, bold=True)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h1.paragraph_format.first_line_indent = Inches(0)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = False
    
    # Heading 2: Left-aligned, Bold, Title Case
    h2 = doc.styles['Heading 2']
    set_style_font(h2, APA_FONT, APA_SIZE, bold=True)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(24)
    h2.paragraph_format.space_after = Pt(12)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h2.paragraph_format.first_line_indent = Inches(0)
    h2.paragraph_format.keep_with_next = True
    
    # Heading 3: Left-aligned, Bold Italic, Title Case
    h3 = doc.styles['Heading 3']
    set_style_font(h3, APA_FONT, APA_SIZE, bold=True, italic=True)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(24)
    h3.paragraph_format.space_after = Pt(12)
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h3.paragraph_format.first_line_indent = Inches(0)
    h3.paragraph_format.keep_with_next = True
    
    # Heading 4: Indented 0.5", Bold, Title Case, Ends with Period
    h4 = doc.styles['Heading 4']
    set_style_font(h4, APA_FONT, APA_SIZE, bold=True)
    h4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h4.paragraph_format.space_before = Pt(24)
    h4.paragraph_format.space_after = Pt(0)
    h4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h4.paragraph_format.first_line_indent = Inches(APA_FIRST_INDENT)
    h4.paragraph_format.keep_with_next = True
    
    # Heading 5: Indented 0.5", Bold Italic, Title Case, Ends with Period
    h5 = doc.styles['Heading 5']
    set_style_font(h5, APA_FONT, APA_SIZE, bold=True, italic=True)
    h5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h5.paragraph_format.space_before = Pt(24)
    h5.paragraph_format.space_after = Pt(0)
    h5.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h5.paragraph_format.first_line_indent = Inches(APA_FIRST_INDENT)
    h5.paragraph_format.keep_with_next = True
    
    # ============================================================
    # TITLE STYLE (for title page - centered, bold, no indent)
    # ============================================================
    title_style = doc.styles['Title']
    set_style_font(title_style, APA_FONT, APA_SIZE, bold=True)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(0)
    title_style.paragraph_format.first_line_indent = Inches(0)
    
    # ============================================================
    # CENTERED STYLE (for title page elements - no indent)
    # ============================================================
    try:
        centered_style = doc.styles.add_style('Centered', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        centered_style = doc.styles['Centered']
    
    set_style_font(centered_style, APA_FONT, APA_SIZE)
    centered_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    centered_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    centered_style.paragraph_format.space_before = Pt(0)
    centered_style.paragraph_format.space_after = Pt(0)
    centered_style.paragraph_format.first_line_indent = Inches(0)
    
    # ============================================================
    # NO INDENT STYLE (for paragraphs that shouldn't have first-line indent)
    # ============================================================
    try:
        no_indent = doc.styles.add_style('No Indent', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        no_indent = doc.styles['No Indent']
    
    set_style_font(no_indent, APA_FONT, APA_SIZE)
    no_indent.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    no_indent.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    no_indent.paragraph_format.space_before = Pt(0)
    no_indent.paragraph_format.space_after = Pt(0)
    no_indent.paragraph_format.first_line_indent = Inches(0)
    
    # ============================================================
    # BLOCK TEXT / QUOTE STYLE (0.5" left indent, no first-line)
    # ============================================================
    try:
        block_style = doc.styles.add_style('Block Text', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        block_style = doc.styles['Block Text']
    
    set_style_font(block_style, APA_FONT, APA_SIZE)
    block_style.paragraph_format.left_indent = Inches(APA_FIRST_INDENT)
    block_style.paragraph_format.first_line_indent = Inches(0)
    block_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    block_style.paragraph_format.space_before = Pt(12)
    block_style.paragraph_format.space_after = Pt(12)
    
    # ============================================================
    # FIRST PARAGRAPH STYLE (no first-line indent after headings)
    # ============================================================
    try:
        first_para = doc.styles.add_style('First Paragraph', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        first_para = doc.styles['First Paragraph']
    
    set_style_font(first_para, APA_FONT, APA_SIZE)
    first_para.paragraph_format.first_line_indent = Inches(0)
    first_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    first_para.paragraph_format.space_before = Pt(0)
    first_para.paragraph_format.space_after = Pt(0)
    
    # ============================================================
    # COMPACT STYLE (for title page - single spacing, no indent)
    # ============================================================
    try:
        compact = doc.styles.add_style('Compact', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        compact = doc.styles['Compact']
    
    set_style_font(compact, APA_FONT, APA_SIZE)
    compact.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    compact.paragraph_format.line_spacing = 1.0
    compact.paragraph_format.space_before = Pt(0)
    compact.paragraph_format.space_after = Pt(0)
    compact.paragraph_format.first_line_indent = Inches(0)
    
    # ============================================================
    # CAPTION STYLE (Figure/Table captions)
    # ============================================================
    caption_style = doc.styles['Caption']
    set_style_font(caption_style, APA_FONT, APA_SIZE)
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    caption_style.paragraph_format.space_before = Pt(12)
    caption_style.paragraph_format.space_after = Pt(12)
    caption_style.paragraph_format.first_line_indent = Inches(0)
    
    # ============================================================
    # TOC STYLES
    # ============================================================
    for i in range(1, 4):
        toc_name = f'TOC {i}'
        try:
            toc_style = doc.styles[toc_name]
            set_style_font(toc_style, APA_FONT, APA_SIZE)
            toc_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            toc_style.paragraph_format.left_indent = Inches(0.25 * (i - 1))
            toc_style.paragraph_format.first_line_indent = Inches(0)
        except KeyError:
            pass
    
    # ============================================================
    # LIST STYLES (bullets and numbers)
    # ============================================================
    list_para = doc.styles['List Paragraph']
    set_style_font(list_para, APA_FONT, APA_SIZE)
    list_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    list_para.paragraph_format.space_before = Pt(0)
    list_para.paragraph_format.space_after = Pt(0)
    
    # ============================================================
    # ADD SAMPLE CONTENT (required for pandoc to inherit styles)
    # This demonstrates each style so pandoc knows they exist
    # ============================================================
    
    # Title page styles
    doc.add_paragraph('Dissertation Title', style='Title')
    doc.add_paragraph('Centered text for title page', style='Centered')
    doc.add_paragraph('No indent paragraph', style='No Indent')
    doc.add_paragraph('Compact style text', style='Compact')
    
    # Body text with first-line indent
    doc.add_paragraph('Body text with first-line indent for APA 7 style formatting.', style='Body Text')
    doc.add_paragraph('Normal paragraph with standard formatting and first-line indent.', style='Normal')
    
    # Headings
    doc.add_heading('Chapter 1: Introduction', level=1)
    doc.add_paragraph('This is the first paragraph after a heading, without first-line indent.', style='First Paragraph')
    doc.add_paragraph('Subsequent paragraphs have the standard 0.5 inch first-line indentation required by APA 7th edition formatting guidelines.')
    
    doc.add_heading('Literature Review', level=2)
    doc.add_paragraph('Level 2 headings are left-aligned and bold.')
    
    doc.add_heading('Theoretical Framework', level=3)
    doc.add_paragraph('Level 3 headings are left-aligned, bold, and italic.')
    
    doc.add_heading('Subsection Heading.', level=4)
    doc.add_paragraph('Level 4 headings are indented 0.5 inches, bold, and end with a period.')
    
    doc.add_heading('Minor Heading.', level=5)
    doc.add_paragraph('Level 5 headings are indented 0.5 inches, bold italic, and end with a period.')
    
    # Block quote
    doc.add_paragraph(
        'This is a block quotation. In APA style, block quotations are used for quotes of 40 words or more. '
        'They are indented 0.5 inches from the left margin and do not use quotation marks.',
        style='Block Text'
    )
    
    # List
    list_para = doc.add_paragraph('Bullet point example', style='List Paragraph')
    
    # ============================================================
    # SAVE REFERENCE DOCUMENT
    # ============================================================
    output_path = Path(__file__).parent.parent / 'templates' / 'reference.docx'
    output_path.parent.mkdir(exist_ok=True)
    doc.save(output_path)
    
    print(f"✓ Created APA 7 reference document: {output_path}")
    print()
    print("APA 7 Styles Configured:")
    print(f"  ├─ Document defaults: {APA_FONT} {APA_SIZE}pt, double-spaced")
    print(f"  ├─ Normal: {APA_FIRST_INDENT}\" first-line indent")
    print(f"  ├─ Body Text: {APA_FIRST_INDENT}\" first-line indent")
    print("  ├─ Title: Centered, Bold, no indent")
    print("  ├─ Centered: Center-aligned, no indent")
    print("  ├─ No Indent: Left-aligned, no first-line indent")
    print("  ├─ Compact: Centered, single-spaced")
    print("  ├─ Heading 1: Centered, Bold (Chapter titles)")
    print("  ├─ Heading 2: Left-aligned, Bold")
    print("  ├─ Heading 3: Left-aligned, Bold Italic")
    print("  ├─ Heading 4: Indented, Bold, ends with period")
    print("  ├─ Heading 5: Indented, Bold Italic, ends with period")
    print("  ├─ Block Text: 0.5\" left indent (for quotes 40+ words)")
    print("  ├─ First Paragraph: No first-line indent")
    print("  ├─ Caption: Figure/table captions")
    print(f"  └─ Page Setup: {APA_MARGIN}\" margins all sides")
    
    return output_path


if __name__ == '__main__':
    create_apa7_reference()
