"""
Format tables in Word documents with professional academic styling.

Usage:
    python format_word_tables.py <input.docx> [output.docx]
    python format_word_tables.py --all   # formats both dissertation and exec summary

Applies:
  - Blue header row (#0078D4) with white bold text
  - Alternating row shading (#F0F0F0 / white)
  - Gray borders (outer #666666, inner #AAAAAA)
  - 10pt header / 9pt body fonts (Calibri)
  - Anti-pagination (cantSplit + keepWithNext)
  - Caption styling above tables
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Color Palette ─────────────────────────────────────────────
HEADER_BG = '0078D4'      # Microsoft blue
HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
ALT_ROW   = 'F0F0F0'      # Light gray for alternating rows
BORDER_OUTER = '666666'
BORDER_INNER = 'AAAAAA'

HEADER_FONT_SIZE = Pt(10)
BODY_FONT_SIZE   = Pt(9)
FONT_NAME = 'Calibri'


def set_cell_shading(cell, color_hex):
    """Apply background shading to a cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tc_pr.append(borders)

    for side, color in [('top', top), ('bottom', bottom),
                        ('left', left), ('right', right)]:
        if color:
            border_el = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" '
                f'w:sz="4" w:space="0" w:color="{color}"/>'
            )
            existing = borders.find(qn(f'w:{side}'))
            if existing is not None:
                borders.remove(existing)
            borders.append(border_el)


def set_row_cant_split(row):
    """Prevent row from splitting across pages."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn('w:cantSplit'))
    if cant_split is None:
        cant_split = parse_xml(f'<w:cantSplit {nsdecls("w")} w:val="true"/>')
        tr_pr.append(cant_split)


def set_row_keep_with_next(row):
    """Keep row with next paragraph/row."""
    for cell in row.cells:
        for para in cell.paragraphs:
            pPr = para._p.get_or_add_pPr()
            kwn = pPr.find(qn('w:keepNext'))
            if kwn is None:
                kwn = parse_xml(f'<w:keepNext {nsdecls("w")} w:val="true"/>')
                pPr.append(kwn)


def set_repeat_header(row):
    """Mark row as repeating table header."""
    tr_pr = row._tr.get_or_add_trPr()
    hdr = tr_pr.find(qn('w:tblHeader'))
    if hdr is None:
        hdr = parse_xml(f'<w:tblHeader {nsdecls("w")} w:val="true"/>')
        tr_pr.append(hdr)


def format_cell_text(cell, font_size, bold=False, color=None):
    """Apply font styling to all runs in a cell."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = font_size
            run.font.name = FONT_NAME
            if bold:
                run.font.bold = True
            if color:
                run.font.color.rgb = color
        # Also set paragraph-level font for empty cells
        if not para.runs and para.text.strip():
            run = para.add_run(para.text)
            para.clear()
            para.add_run(run.text)


def set_table_width(table):
    """Set table to auto-fit window width."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tbl_pr)
    # Remove existing width
    existing_w = tbl_pr.find(qn('w:tblW'))
    if existing_w is not None:
        tbl_pr.remove(existing_w)
    # Set to 100% page width
    tbl_w = parse_xml(
        f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>'
    )
    tbl_pr.append(tbl_w)


def format_table(table, table_index):
    """Apply professional academic styling to a single table."""
    rows = table.rows
    if not rows:
        return

    set_table_width(table)

    for row_idx, row in enumerate(rows):
        is_header = (row_idx == 0)

        set_row_cant_split(row)
        if row_idx < len(rows) - 1:
            set_row_keep_with_next(row)

        if is_header:
            set_repeat_header(row)

        for col_idx, cell in enumerate(row.cells):
            num_cols = len(row.cells)
            # Determine border colors
            is_top = (row_idx == 0)
            is_bottom = (row_idx == len(rows) - 1)
            is_left = (col_idx == 0)
            is_right = (col_idx == num_cols - 1)

            set_cell_borders(
                cell,
                top=BORDER_OUTER if is_top else BORDER_INNER,
                bottom=BORDER_OUTER if is_bottom else BORDER_INNER,
                left=BORDER_OUTER if is_left else BORDER_INNER,
                right=BORDER_OUTER if is_right else BORDER_INNER,
            )

            if is_header:
                set_cell_shading(cell, HEADER_BG)
                format_cell_text(cell, HEADER_FONT_SIZE, bold=True,
                                 color=HEADER_FG)
            else:
                if row_idx % 2 == 0:
                    set_cell_shading(cell, ALT_ROW)
                format_cell_text(cell, BODY_FONT_SIZE)


def style_table_captions(doc):
    """Style paragraphs that look like table captions (bold, keep with next)."""
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith('Table ') and len(text) < 200:
            # Looks like a table caption — style it
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_NAME
            # Keep with next to prevent orphan captions
            pPr = para._p.get_or_add_pPr()
            kwn = pPr.find(qn('w:keepNext'))
            if kwn is None:
                kwn = parse_xml(f'<w:keepNext {nsdecls("w")} w:val="true"/>')
                pPr.append(kwn)


def format_document(input_path: str, output_path: str = None):
    """Format all tables in a Word document."""
    input_path = Path(input_path)
    if output_path is None:
        output_path = input_path
    else:
        output_path = Path(output_path)

    doc = Document(str(input_path))
    table_count = len(doc.tables)
    print(f'  Found {table_count} tables in {input_path.name}')

    for i, table in enumerate(doc.tables):
        format_table(table, i)

    style_table_captions(doc)

    doc.save(str(output_path))
    size_kb = output_path.stat().st_size / 1024
    print(f'  Saved: {output_path} ({size_kb:.0f} KB)')
    return table_count


def main():
    if len(sys.argv) < 2:
        print('Usage: python format_word_tables.py <input.docx> [output.docx]')
        print('       python format_word_tables.py --all')
        return 1

    if sys.argv[1] == '--all':
        print('Word Table Formatter — Batch Mode')
        print('=' * 50)

        script_dir = Path(__file__).parent
        repo_root = script_dir.parent
        files = [
            (script_dir / 'output' / 'AIRS_Dissertation.docx', None),
            (repo_root / 'defence' / 'exports' / 'AIRS-Executive-Summary.docx', None),
        ]

        total = 0
        for full, out_path in files:
            if not full.exists():
                print(f'  SKIP  {full.name} (not found)')
                continue
            out = full if out_path is None else Path(out_path)
            total += format_document(str(full), str(out))

        print(f'\nTotal: {total} tables formatted')
        return 0

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    print('Word Table Formatter')
    print('=' * 50)
    format_document(input_file, output_file)
    return 0


if __name__ == '__main__':
    sys.exit(main())
